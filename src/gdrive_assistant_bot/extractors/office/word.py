from __future__ import annotations

import os
import subprocess
import tempfile
from typing import Any

from docx import Document as DocxDocument

from ..base import ExtractedContent, ExtractionContext, FileExtractor

DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_DOC_MIME_TYPE = "application/msword"


class DocxExtractor(FileExtractor):
    """Extract text from DOCX files."""

    @property
    def mime_types(self) -> list[str]:
        return [DOCX_MIME_TYPE]

    @property
    def file_extensions(self) -> list[str]:
        return ["docx"]

    def can_extract(self, file_meta: dict[str, Any]) -> bool:
        if file_meta.get("mimeType") == DOCX_MIME_TYPE:
            return True
        return self._extension(file_meta) == "docx"

    def extract(self, file_meta: dict[str, Any], context: ExtractionContext) -> ExtractedContent:
        size = self._to_int(file_meta.get("size"))
        max_bytes = int(context.settings.OFFICE_MAX_FILE_SIZE_MB * 1024 * 1024)
        if size and size > max_bytes:
            return ExtractedContent(
                text="", file_type="docx", metadata={"skipped": "size_limit", "size_bytes": size}
            )

        docx_bytes = context.download_binary(file_meta["id"])
        text = self._extract_docx(docx_bytes)
        return ExtractedContent(
            text=text.strip(),
            file_type="docx",
            metadata={"mime_type": file_meta.get("mimeType"), "file_size_bytes": len(docx_bytes)},
        )

    @staticmethod
    def _extract_docx(docx_bytes: bytes) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name

        try:
            doc = DocxDocument(tmp_path)
            lines: list[str] = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    lines.append(text)

            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)
        finally:
            os.unlink(tmp_path)

    @staticmethod
    def _extension(file_meta: dict[str, Any]) -> str | None:
        ext = file_meta.get("fileExtension")
        if isinstance(ext, str) and ext.strip():
            return ext.lower().lstrip(".")

        name = file_meta.get("name")
        if not isinstance(name, str) or "." not in name:
            return None
        return name.rsplit(".", 1)[-1].strip().lower() or None

    @staticmethod
    def _to_int(value: Any) -> int | None:
        if isinstance(value, int):
            return value
        if isinstance(value, str) and value.isdigit():
            return int(value)
        return None


class DocExtractor(FileExtractor):
    """Extract text from legacy DOC files."""

    @property
    def mime_types(self) -> list[str]:
        return [_DOC_MIME_TYPE]

    @property
    def file_extensions(self) -> list[str]:
        return ["doc"]

    def can_extract(self, file_meta: dict[str, Any]) -> bool:
        if file_meta.get("mimeType") == _DOC_MIME_TYPE:
            return True
        return self._extension(file_meta) == "doc"

    def extract(self, file_meta: dict[str, Any], context: ExtractionContext) -> ExtractedContent:
        size = self._to_int(file_meta.get("size"))
        max_bytes = int(context.settings.OFFICE_MAX_FILE_SIZE_MB * 1024 * 1024)
        if size and size > max_bytes:
            return ExtractedContent(
                text="", file_type="doc", metadata={"skipped": "size_limit", "size_bytes": size}
            )

        doc_bytes = context.download_binary(file_meta["id"])
        text = self._extract_doc(doc_bytes)
        return ExtractedContent(
            text=text.strip(),
            file_type="doc",
            metadata={"mime_type": file_meta.get("mimeType"), "file_size_bytes": len(doc_bytes)},
        )

    @staticmethod
    def _extract_doc(doc_bytes: bytes) -> str:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as tmp:
            tmp.write(doc_bytes)
            tmp_path = tmp.name

        try:
            result = subprocess.run(["catdoc", tmp_path], capture_output=True, check=False)
        except FileNotFoundError as exc:
            raise RuntimeError("Legacy DOC extraction requires the 'catdoc' binary.") from exc
        finally:
            os.unlink(tmp_path)

        if result.returncode != 0:
            stderr = (result.stderr or b"").decode("utf-8", errors="replace").strip()
            raise RuntimeError(f"catdoc failed to extract DOC file: {stderr or 'unknown error'}")

        return (result.stdout or b"").decode("utf-8", errors="replace")

    @staticmethod
    def _extension(file_meta: dict[str, Any]) -> str | None:
        ext = file_meta.get("fileExtension")
        if isinstance(ext, str) and ext.strip():
            return ext.lower().lstrip(".")

        name = file_meta.get("name")
        if not isinstance(name, str) or "." not in name:
            return None
        return name.rsplit(".", 1)[-1].strip().lower() or None

    @staticmethod
    def _to_int(value: Any) -> int | None:
        if isinstance(value, int):
            return value
        if isinstance(value, str) and value.isdigit():
            return int(value)
        return None
